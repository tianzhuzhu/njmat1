import os
import datetime
import pandas as pd
from docx.oxml.ns import qn
from docx.shared import Pt,RGBColor

from docx import Document
from docx.shared import Inches
def citymorethansevendays(moretsc):
    print(moretsc)
    list = moretsc.columns.tolist()
    copylist = list.copy()
    copylist.remove('部门')
    copylist.remove('汇总')
    str1=''
    count = 0
    lastcount = 0
    for index, row in moretsc.iterrows():
        count = count + 1
        if (row['部门'] != '总计'):
            str1 += row['部门'] + str(row['汇总']) + '件('
            for v in copylist:
                if (str(row[v]).lower() != 'nan'):
                    str1 += v + str(row[v]) + '件、'
            str1 = str1[0:-1]
            str1 += ')、'
        if (count >= 3 and lastcount != row['汇总']):
            break
        lastcount = row['汇总']
    str1 = str1[0:-1]
    print(str1)
    print('-----------市政工程完成多于七天完成-----------')
    return str1
def last_day_of_month(any_day):
    """
    获取获得一个月中的最后一天
    :param any_day: 任意日期
    :return: string
    """
    next_month = any_day.replace(day=28) + datetime.timedelta(days=4)  # this will never fail
    return next_month - datetime.timedelta(days=next_month.day)

# 注意: 年月日，这些变量必须是数字，否则报错！
year = 2019 # 年
month = 5  # 月
day = 16 # 日
def removeUnameColumns(data):
    data.dropna(inplace=True,how='all')
    columns=data.columns.tolist()
    for i in columns:
        if(i.startswith('Unnamed')):
            columns.remove(i)
    return pd.DataFrame(data,columns=columns)
def getOutPutName(file):

    l = file.rfind('\\')
    r = file.rfind('月')
    path=file[:l]
    mo = file[l+1:r]
    path=os.path.join(path,'result')
    return path

def others(data):
    str1=''
    count = 0
    lastcount = 0
    for index, row in data.iterrows():
        count = count + 1
        if (row['部门'] != '总计'):
            str1 += row['部门'] + row['汇总'] + '件、'
        if (count >= 3 and lastcount != row['汇总']):
            break
        lastcount = row['汇总']
    str1 = str1[0:-1]
    print(str1)
    print('-----------xx完成-----------')
    return str1

def readexcel(file):
    totalcount=pd.read_excel(file,sheet_name='总清单汇总',dtype=str)
    moretsc=pd.read_excel(file,sheet_name='超时七天汇总表',dtype=str)
    moretsc = removeUnameColumns(moretsc)
    lesstsc=pd.read_excel(file,sheet_name='未超时七天汇总表',dtype=str)
    total=totalcount.loc[totalcount['部门']=='总计',['汇总']]
    total=total.iloc[0,0]
    print(str(total))

    print('----------total----------')
    print(total)
    numofmtsd=moretsc.loc[moretsc['部门']=='总计','汇总']
    numofmtsd=numofmtsd.iloc[0]
    numofltsd=lesstsc.loc[lesstsc['部门']=='总计','汇总']
    numofltsd = numofltsd.iloc[0]
    #整治工单
    str1=''
    if ( file.rfind('市政工程')<file.rfind('整治工单')):
        str1=others(moretsc)
    #市政工程
    elif ( file.rfind('整治工单')<file.rfind('市政工程')):
        str1=citymorethansevendays(moretsc)
    print(lesstsc)
    str2 = others(lesstsc)

    return total,numofmtsd,numofltsd,str1,str2


if __name__=="__main__":
    print('------------产生word文件---------------')

    file1=input('请输入市政工程位置：')
    file2 = input('请输入整治工程位置：')
    tomonth = datetime.datetime.now().month
    today=datetime.datetime.now().day
    print(today)
    lastday=last_day_of_month(datetime.datetime.now()).day
    document=Document()
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Normal'].font.size = Pt(10.5)
    document.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
    pf1='各位领导、同事：'
    pf2='  截止{}月{}日在途市政工程单共{}件。在途超时7天工单{}件，其中{}；在途未超7天工单{}件，其中{}。'
    total,numofmtsd,numofltsd,str1,str2=readexcel(file1)
    print('------------------------')
    print(total)
    pf2=pf2.format(tomonth,day,total,numofmtsd,str1,numofltsd,str2)
    print(pf2)
    pf3='  截止{}月{}日在途整治工单共{}件。在途超时7天工单{}件，其中{}；在途未超7天工单{}件，其中{}，详情见附件。'
    total, numofmtsd, numofltsd, str1, str2 = readexcel(file2)

    pf3 = pf3.format(tomonth, day, total, numofmtsd, str1, numofltsd, str2)
    print(pf3)
    pf4='  现将要求通知如下：'
    pf5=' 一、在途超时7天工单'
    pf6='{0}月{1}日前联系用户核实、处理、归档。'
    pf6=pf6.format(str(tomonth),lastday)
    pf7=' 二、在途未超7天工单'
    pf8='1、{0}月{1}日下班前联系用户核实、处理、归档。如有问题，请反馈原因。'
    pf8=pf8.format(str(tomonth),lastday)
    pf9='2、对拆迁地、施工工地、外力影响等不可抗力因素导致暂时无法修复的，上传照片，每周联系用户，反馈有效处理进展。'
    pf10='整治工单和长时间整治不到位或整治时间超过30天，且涉嫌随意回单的，按工单量考核到部门，≤10'
    pf11='单考核50元，每增加10单追加考核50元。'
    pf12='请各部门及时做好在途市政工程单和整治工单的清理工作，谢谢！'


    paragraph1 = document.add_paragraph(pf1)
    paragraph2 = document.add_paragraph(pf2)
    paragraph3 = document.add_paragraph(pf3)
    paragraph4 = document.add_paragraph(pf4)
    paragraph5 = document.add_paragraph(pf5)
    paragraph6 = document.add_paragraph(pf6)
    paragraph7 = document.add_paragraph(pf7)
    paragraph8 = document.add_paragraph(pf8)
    paragraph9 = document.add_paragraph(pf9)
    paragraph10 = document.add_paragraph(pf10)
    paragraph11 = document.add_paragraph(pf11)
    paragraph12 = document.add_paragraph(pf12)
    document.save(str(tomonth)+'月'+str(today)+'日'+"市政工程整治工单通报.docx")
